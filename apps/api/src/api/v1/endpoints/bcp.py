"""
BCP Generator API — generate Business Continuity Plans via NVIDIA LLM.

Endpoints:
- POST /bcp/generate — generate BCP from entity, scenario, jurisdiction, capabilities
- GET /bcp/config — return sector and scenario config for UI
- POST /bcp/export/pdf — export generated BCP content as PDF
- POST /bcp/export/word — export generated BCP content as Word (DOCX)
"""
from __future__ import annotations

import io
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from src.core.security import get_current_user_optional
from src.models.user import User
from src.services.bcp_config import (
    SECTOR_BCP_CONFIG,
    SCENARIO_BCP_SPECIFICS,
    get_sector_bcp_config,
    get_scenario_bcp_specifics,
)
from src.services.bcp_prompts import BCP_SYSTEM_PROMPT, build_bcp_user_prompt
from src.services.nvidia_llm import NVIDIALLMService
from src.services.nvidia_llm import LLMModel
from src.services.pdf_report import generate_bcp_pdf, HAS_PDF

logger = logging.getLogger(__name__)

router = APIRouter()


# =============================================================================
# REQUEST / RESPONSE MODELS
# =============================================================================


class EntityLocation(BaseModel):
    """Entity location."""
    city: Optional[str] = None
    country: Optional[str] = None
    region: Optional[str] = None


class BCPEntity(BaseModel):
    """Organization/entity for BCP."""
    name: str = Field(..., min_length=1)
    type: str = Field(..., description="Sector: insurance, healthcare, financial, etc.")
    subtype: Optional[str] = None
    location: Optional[EntityLocation] = None
    size: str = Field("medium", pattern="^(large|medium|small)$")
    employees: Optional[int] = Field(None, ge=0)
    critical_functions: Optional[List[str]] = None
    dependencies: Optional[List[str]] = None


class BCPScenario(BaseModel):
    """Threat scenario for BCP."""
    type: str = Field(..., description="flood, cyber, pandemic, etc.")
    severity: float = Field(0.5, ge=0, le=1)
    duration_estimate: Optional[str] = None
    specific_threat: Optional[str] = None


class BCPJurisdiction(BaseModel):
    """Regulatory jurisdiction."""
    primary: str = Field("EU", description="EU, Germany, USA, UK, Japan, etc.")
    secondary: Optional[List[str]] = None
    regulations: Optional[List[str]] = None


class BCPExistingCapabilities(BaseModel):
    """Existing BCP capabilities."""
    has_bcp: bool = False
    last_test_date: Optional[str] = None
    backup_site: bool = False
    remote_work_ready: bool = False


class BCPGenerateRequest(BaseModel):
    """Request body for BCP generation."""
    entity: BCPEntity
    scenario: BCPScenario
    jurisdiction: BCPJurisdiction = Field(default_factory=BCPJurisdiction)
    existing_capabilities: BCPExistingCapabilities = Field(default_factory=BCPExistingCapabilities)


class BCPGenerateResponse(BaseModel):
    """Response from BCP generation."""
    content: str
    model: str = ""
    tokens_used: Optional[int] = None
    finish_reason: Optional[str] = None


# =============================================================================
# LLM SERVICE
# =============================================================================

_llm_service: Optional[NVIDIALLMService] = None


def get_llm_service() -> NVIDIALLMService:
    global _llm_service
    if _llm_service is None:
        _llm_service = NVIDIALLMService()
    return _llm_service


# =============================================================================
# ENDPOINTS
# =============================================================================


@router.post("/generate", response_model=BCPGenerateResponse)
async def generate_bcp(
    request: BCPGenerateRequest,
    current_user: Optional[User] = Depends(get_current_user_optional),
) -> BCPGenerateResponse:
    """
    Generate a Business Continuity Plan using NVIDIA LLM.

    Request: entity (name, sector, location, size, ...), scenario (type, severity, ...),
    jurisdiction (primary, regulations), existing_capabilities.
    Returns full BCP text.
    """
    llm = get_llm_service()
    if not llm.is_available:
        raise HTTPException(
            status_code=503,
            detail="LLM service not available (NVIDIA API key or NIM not configured)",
        )

    sector_key = (request.entity.type or "enterprise").strip().lower().replace("-", "_")
    scenario_key = (request.scenario.type or "flood").strip().lower().replace(" ", "_")
    sector_config = get_sector_bcp_config(sector_key) or get_sector_bcp_config("enterprise")
    scenario_config = get_scenario_bcp_specifics(scenario_key) or get_scenario_bcp_specifics("flood")

    entity_dict: Dict[str, Any] = {
        "name": request.entity.name,
        "type": sector_key,
        "subtype": request.entity.subtype,
        "location": request.entity.location.model_dump() if request.entity.location else {},
        "size": request.entity.size,
        "employees": request.entity.employees,
        "critical_functions": request.entity.critical_functions or [],
        "dependencies": request.entity.dependencies or [],
    }
    scenario_dict: Dict[str, Any] = {
        "type": scenario_key,
        "severity": request.scenario.severity,
        "duration_estimate": request.scenario.duration_estimate or "",
        "specific_threat": request.scenario.specific_threat or "",
    }
    jurisdiction_dict: Dict[str, Any] = {
        "primary": request.jurisdiction.primary or "EU",
        "secondary": request.jurisdiction.secondary or [],
        "regulations": request.jurisdiction.regulations,
    }
    capabilities_dict: Dict[str, Any] = {
        "has_bcp": request.existing_capabilities.has_bcp,
        "last_test_date": request.existing_capabilities.last_test_date,
        "backup_site": request.existing_capabilities.backup_site,
        "remote_work_ready": request.existing_capabilities.remote_work_ready,
    }

    user_prompt = build_bcp_user_prompt(
        entity_dict,
        scenario_dict,
        jurisdiction_dict,
        capabilities_dict,
        sector_config,
        scenario_config,
    )

    try:
        response = await llm.generate(
            prompt=user_prompt,
            model=LLMModel.LLAMA_70B,
            max_tokens=8000,
            temperature=0.3,
            system_prompt=BCP_SYSTEM_PROMPT,
        )
    except Exception as e:
        logger.exception("BCP LLM generation failed: %s", e)
        raise HTTPException(status_code=500, detail=f"BCP generation failed: {str(e)}") from e

    return BCPGenerateResponse(
        content=response.content or "",
        model=response.model or "meta/llama-3.1-70b-instruct",
        tokens_used=getattr(response, "tokens_used", None),
        finish_reason=getattr(response, "finish_reason", None),
    )


@router.get("/config")
async def get_bcp_config(
    current_user: Optional[User] = Depends(get_current_user_optional),
) -> Dict[str, Any]:
    """
    Return sector and scenario BCP config for the UI.

    Sectors: critical_functions, regulations by jurisdiction, key_roles.
    Scenarios: activation_criteria, immediate_actions, resources_needed, special_considerations.
    """
    return {
        "sectors": SECTOR_BCP_CONFIG,
        "scenarios": SCENARIO_BCP_SPECIFICS,
    }


# =============================================================================
# EXPORT (PDF / WORD)
# =============================================================================


class BCPExportRequest(BaseModel):
    """Request body for BCP export (PDF or Word)."""
    content: str = Field(..., min_length=1)


def _bcp_pdf_response(data: bytes, filename: str) -> StreamingResponse:
    """Return PDF as downloadable response."""
    return StreamingResponse(
        io.BytesIO(data),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Type": "application/pdf",
        },
    )


def _bcp_docx_response(data: bytes, filename: str) -> StreamingResponse:
    """Return DOCX as downloadable response."""
    return StreamingResponse(
        io.BytesIO(data),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
    )


@router.post("/export/pdf")
async def export_bcp_pdf(
    request: BCPExportRequest,
    current_user: Optional[User] = Depends(get_current_user_optional),
) -> StreamingResponse:
    """Export BCP content as PDF. Requires ReportLab."""
    if not HAS_PDF:
        raise HTTPException(
            status_code=503,
            detail="PDF generation not available. Install reportlab.",
        )
    try:
        pdf_bytes = generate_bcp_pdf(request.content)
    except Exception as e:
        logger.exception("BCP PDF export failed: %s", e)
        raise HTTPException(status_code=500, detail=f"PDF export failed: {str(e)}") from e
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    return _bcp_pdf_response(pdf_bytes, f"bcp_{ts}.pdf")


@router.post("/export/word")
async def export_bcp_word(
    request: BCPExportRequest,
    current_user: Optional[User] = Depends(get_current_user_optional),
) -> StreamingResponse:
    """Export BCP content as Word (DOCX). Requires python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="Word export not available. Install python-docx.",
        )
    try:
        doc = Document()
        doc.add_heading("Business Continuity Plan", 0)
        lines = (request.content or "").split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            i += 1
            if not stripped:
                doc.add_paragraph()
                continue
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:].strip(), level=1)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:].strip(), level=2)
            else:
                doc.add_paragraph(stripped)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        return _bcp_docx_response(buf.getvalue(), f"bcp_{ts}.docx")
    except Exception as e:
        logger.exception("BCP Word export failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Word export failed: {str(e)}") from e
